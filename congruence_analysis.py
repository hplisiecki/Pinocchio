# congruence_analysis.py
#
# Tests the self-model hypothesis by computing Tucker's congruence coefficient (phi)
# between loading matrices across the three prompting conditions for each scale.
#
# Predicted ordering under the self-model hypothesis:
#   phi(llm_analog, neutral) > phi(human_simulation, llm_analog) ≈ phi(human_simulation, neutral)
#
# llm_analog and neutral are both "respond as yourself" conditions — the functional-analog
# framing in llm_analog is designed to reduce refusals, not to change the self-model.
# human_simulation imposes an external human self-model, producing a categorically different
# response structure. The two respond-as-self conditions should therefore be most similar
# to each other, and human_simulation should be roughly equally distant from both.
#
# Usage:
#   conda run -n llmlatent python congruence_analysis.py

import os
import re
import numpy as np
import pandas as pd
from scipy.optimize import linear_sum_assignment
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EFA_DIR     = "data/efa_results"
OUTPUT_DOC  = "data/congruence_report.docx"
OUTPUT_CSV  = "data/congruence_results.csv"
CONDITIONS  = ["llm_analog", "neutral", "human_simulation"]
PAIRS = [
    ("human_simulation", "llm_analog"),   # both have self-models — predicted highest
    ("llm_analog",       "neutral"),       # one has self-model   — predicted middle
    ("human_simulation", "neutral"),       # most different        — predicted lowest
]

# ---------------------------------------------------------------------------
# Tucker's congruence
# ---------------------------------------------------------------------------

def tucker_phi(v1: np.ndarray, v2: np.ndarray) -> float:
    denom = np.sqrt(np.sum(v1 ** 2) * np.sum(v2 ** 2))
    return float(np.sum(v1 * v2) / denom) if denom > 0 else 0.0


def mean_congruence(A: np.ndarray, B: np.ndarray) -> float:
    """
    Match factors between A and B using the Hungarian algorithm (maximising
    absolute congruence, allowing factor reflection), then return the mean
    Tucker's phi across matched pairs.
    """
    nA, nB = A.shape[1], B.shape[1]
    cost = np.zeros((nA, nB))
    for i in range(nA):
        for j in range(nB):
            cost[i, j] = -abs(tucker_phi(A[:, i], B[:, j]))
    rows, cols = linear_sum_assignment(cost)
    phis = [abs(tucker_phi(A[:, r], B[:, c])) for r, c in zip(rows, cols)]
    return float(np.mean(phis))


# ---------------------------------------------------------------------------
# Load a loading matrix for one scale × condition
# ---------------------------------------------------------------------------

def load_matrix(quest: str, condition: str):
    """
    Returns (items: list, matrix: np.ndarray) or None if file not found.
    Items are the row labels; matrix columns are factors/components.
    """
    safe = re.sub(r'[\\/*?:"<>|]', "_", quest)
    path = os.path.join(EFA_DIR, f"{safe}__{condition}.csv")
    if not os.path.exists(path):
        return None
    df = pd.read_csv(path, index_col=0)
    return list(df.index), df.values


# ---------------------------------------------------------------------------
# Main analysis
# ---------------------------------------------------------------------------

def run_analysis():
    summary_df = pd.read_csv(os.path.join(EFA_DIR, "summary.csv"))
    questionnaires = sorted(summary_df["questionnaire"].unique())

    rows = []

    for quest in questionnaires:
        # Load all three matrices
        matrices = {}
        for cond in CONDITIONS:
            result = load_matrix(quest, cond)
            if result is not None:
                matrices[cond] = result  # (items, array)

        if len(matrices) < 3:
            print(f"  SKIP {quest}: missing condition(s)")
            continue

        # Find items present in all three conditions (some may have been dropped
        # for zero variance in one condition but not another)
        item_sets = [set(matrices[c][0]) for c in CONDITIONS]
        common_items = sorted(item_sets[0].intersection(*item_sets[1:]))

        if len(common_items) < 3:
            print(f"  SKIP {quest}: fewer than 3 common items across conditions")
            continue

        # Subset each matrix to common items
        trimmed = {}
        for cond in CONDITIONS:
            items, mat = matrices[cond]
            idx = [items.index(it) for it in common_items]
            trimmed[cond] = mat[idx, :]

        # Compute pairwise congruence
        pair_phis = {}
        for c1, c2 in PAIRS:
            pair_phis[(c1, c2)] = mean_congruence(trimmed[c1], trimmed[c2])

        phi_hs_la = pair_phis[("human_simulation", "llm_analog")]
        phi_la_n  = pair_phis[("llm_analog",       "neutral")]
        phi_hs_n  = pair_phis[("human_simulation", "neutral")]

        # Does the predicted ordering hold?
        # phi(la, n) > phi(hs, la) AND phi(la, n) > phi(hs, n)
        # i.e. the two respond-as-self conditions are more similar to each other
        # than either is to human_simulation
        ordering_holds = (phi_la_n > phi_hs_la) and (phi_la_n > phi_hs_n)

        # Number of factors per condition
        nf = {c: trimmed[c].shape[1] for c in CONDITIONS}

        rows.append({
            "questionnaire":         quest,
            "n_factors_llm_analog":  nf["llm_analog"],
            "n_factors_neutral":     nf["neutral"],
            "n_factors_human_sim":   nf["human_simulation"],
            "phi_hs_la":             round(phi_hs_la, 3),
            "phi_la_n":              round(phi_la_n,  3),
            "phi_hs_n":              round(phi_hs_n,  3),
            "ordering_holds":        ordering_holds,
        })
        status = "OK" if ordering_holds else "--"
        print(f"  [{status}] {quest}: hs-la={phi_hs_la:.3f}  la-n={phi_la_n:.3f}  hs-n={phi_hs_n:.3f}")

    results = pd.DataFrame(rows)
    results.to_csv(OUTPUT_CSV, index=False)

    n_holds = results["ordering_holds"].sum()
    n_total = len(results)
    print(f"\nOrdering holds for {n_holds}/{n_total} scales ({100*n_holds/n_total:.1f}%)")
    print(f"\nMean phi values:")
    print(f"  phi(hs, la) = {results['phi_hs_la'].mean():.3f}  (both have self-model)")
    print(f"  phi(la,  n) = {results['phi_la_n'].mean():.3f}  (one has self-model)")
    print(f"  phi(hs,  n) = {results['phi_hs_n'].mean():.3f}  (neither/most different)")

    return results


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, size=9, bold=False, center=False):
    cell.text = str(text)
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(size)
    run.bold = bold
    if center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_report(results: pd.DataFrame):
    doc = Document()

    title = doc.add_heading("Self-Model Hypothesis: Cross-Condition Congruence Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Explanation ---
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This report tests the self-model hypothesis using Tucker's congruence coefficient (phi), "
        "a standard psychometric measure of factor structure similarity (range 0–1; "
        "phi > 0.95 = near-identical, phi > 0.85 = similar, phi < 0.70 = dissimilar). "
        "For each scale, phi is computed between each pair of prompting conditions after "
        "aligning factors via the Hungarian algorithm."
    )
    doc.add_paragraph(
        "Predicted ordering under the self-model hypothesis: "
        "phi(llm_analog, neutral) > phi(human_simulation, llm_analog) ≈ phi(human_simulation, neutral). "
        "Both llm_analog and neutral are 'respond as yourself' conditions — the functional-analog "
        "framing in llm_analog is designed to reduce refusals, not to impose a different self-model. "
        "human_simulation imposes an external human self-model, producing categorically different "
        "responding. The two respond-as-self conditions should be most similar to each other, and "
        "human_simulation should be roughly equally distant from both."
    )

    # --- Summary statistics ---
    doc.add_heading("Summary", level=1)
    n_holds  = results["ordering_holds"].sum()
    n_total  = len(results)
    pct      = 100 * n_holds / n_total

    p = doc.add_paragraph()
    p.add_run(f"Predicted ordering holds for {n_holds} / {n_total} scales ({pct:.1f}%).\n").bold = True

    means = {
        "phi(human_simulation, llm_analog)": results["phi_hs_la"].mean(),
        "phi(llm_analog, neutral)":          results["phi_la_n"].mean(),
        "phi(human_simulation, neutral)":    results["phi_hs_n"].mean(),
    }
    for label, val in means.items():
        doc.add_paragraph(f"  Mean {label} = {val:.3f}", style="List Bullet")

    predicted = (means["phi(llm_analog, neutral)"] > means["phi(human_simulation, llm_analog)"] and
                 means["phi(llm_analog, neutral)"] > means["phi(human_simulation, neutral)"])
    hs_symmetric = abs(means["phi(human_simulation, llm_analog)"] -
                       means["phi(human_simulation, neutral)"]) < 0.05
    doc.add_paragraph(
        f"The mean phi values follow the predicted ordering: {'YES' if predicted else 'NO'} "
        f"(phi(la,n)={means['phi(llm_analog, neutral)']:.3f} > "
        f"phi(hs,la)={means['phi(human_simulation, llm_analog)']:.3f} ≈ "
        f"phi(hs,n)={means['phi(human_simulation, neutral)']:.3f}). "
        f"human_simulation is symmetrically distant from both respond-as-self conditions: "
        f"{'YES' if hs_symmetric else 'NO'} "
        f"(difference = {abs(means['phi(human_simulation, llm_analog)'] - means['phi(human_simulation, neutral)']):.3f})."
    )

    # --- Per-scale table ---
    doc.add_heading("Per-Scale Results", level=1)
    doc.add_paragraph(
        "Green = predicted ordering holds (phi(la,n) is the highest of the three pairs). "
        "Phi values are colour-coded: dark green >= 0.90, light green >= 0.70, "
        "yellow >= 0.50, red < 0.50."
    )

    headers = [
        "Questionnaire",
        "F (LA)", "F (N)", "F (HS)",
        "phi(HS,LA)", "phi(LA,N)", "phi(HS,N)",
        "Ordering"
    ]
    table = doc.add_table(rows=1 + len(results), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell_text(table.cell(0, i), h, bold=True, center=(i > 0))
        set_cell_bg(table.cell(0, i), "D9E1F2")

    def phi_color(v):
        if v >= 0.90: return "C6EFCE"
        if v >= 0.70: return "E2EFDA"
        if v >= 0.50: return "FFEB9C"
        return "FFC7CE"

    for ri, (_, row) in enumerate(results.iterrows(), start=1):
        vals = [
            row["questionnaire"],
            row["n_factors_llm_analog"],
            row["n_factors_neutral"],
            row["n_factors_human_sim"],
            row["phi_hs_la"],
            row["phi_la_n"],
            row["phi_hs_n"],
            "YES" if row["ordering_holds"] else "NO",
        ]
        for ci, v in enumerate(vals):
            cell_text(table.cell(ri, ci), v, center=(ci > 0))

        for ci, phi_val in [(4, row["phi_hs_la"]),
                            (5, row["phi_la_n"]),
                            (6, row["phi_hs_n"])]:
            set_cell_bg(table.cell(ri, ci), phi_color(phi_val))

        row_color = "E2EFDA" if row["ordering_holds"] else "FFF2CC"
        set_cell_bg(table.cell(ri, 7), row_color)
        if ri % 2 == 0:
            for ci in [0]:
                set_cell_bg(table.cell(ri, ci), "F2F2F2")

    widths = [Inches(2.5)] + [Inches(0.45)] * 3 + [Inches(0.75)] * 3 + [Inches(0.6)]
    for row in table.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = w

    doc.save(OUTPUT_DOC)
    print(f"\nReport saved to {OUTPUT_DOC}")
    print(f"Results CSV saved to {OUTPUT_CSV}")


if __name__ == "__main__":
    results = run_analysis()
    build_report(results)
